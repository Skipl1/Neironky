#!/usr/bin/env python3
"""
DOCX Content Extractor for Typst Conversion
Extracts: paragraphs, mathematical formulas (OMML -> LaTeX), and images
"""

import os
import sys
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Fix Windows console encoding
if sys.platform == 'win32':
    import codecs
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

# Constants
DOCX_PATH = r"D:\Для Windows 11\Всё подряд\Users\kreck\Desktop\Прочее\Методы оптимизации\Отчёты\Optimization-Methods\lab2\Привалихин Дмитрий_1 часть.docx"
OUTPUT_DIR = r"D:\Для Windows 11\Всё подряд\Users\kreck\Desktop\Прочее\Методы оптимизации\Отчёты\Optimization-Methods\lab2\extracted_content"

# OMML to LaTeX conversion mappings
OMML_TO_LATEX = {
    'frac': lambda elem: f"\\frac{{{get_omml_text(elem.find('.//{n}num'))}}}{{{get_omml_text(elem.find('.//{n}den'))}}}",
    'sqrt': lambda elem: f"\\sqrt{{{get_omml_text(elem)}}}",
    'sup': lambda elem: f"^{get_omml_text(elem)}",
    'sub': lambda elem: f"_{{{get_omml_text(elem)}}}",
    'sum': lambda elem: "\\sum",
    'prod': lambda elem: "\\prod",
    'int': lambda elem: "\\int",
    'infin': lambda elem: "\\infty",
    'ne': lambda elem: "\\neq",
    'leq': lambda elem: "\\leq",
    'geq': lambda elem: "\\geq",
    'cdot': lambda elem: "\\cdot",
    'times': lambda elem: "\\times",
    'div': lambda elem: "\\div",
    'pm': lambda elem: "\\pm",
    'alpha': lambda elem: "\\alpha",
    'beta': lambda elem: "\\beta",
    'gamma': lambda elem: "\\gamma",
    'delta': lambda elem: "\\delta",
    'epsilon': lambda elem: "\\epsilon",
    'theta': lambda elem: "\\theta",
    'lambda': lambda elem: "\\lambda",
    'pi': lambda elem: "\\pi",
    'sigma': lambda elem: "\\sigma",
    'phi': lambda elem: "\\phi",
    'omega': lambda elem: "\\omega",
    'Delta': lambda elem: "\\Delta",
    'nabla': lambda elem: "\\nabla",
    'partial': lambda elem: "\\partial",
    'forall': lambda elem: "\\forall",
    'exists': lambda elem: "\\exists",
    'in': lambda elem: "\\in",
    'notin': lambda elem: "\\notin",
    'subset': lambda elem: "\\subset",
    'subseteq': lambda elem: "\\subseteq",
    'cup': lambda elem: "\\cup",
    'cap': lambda elem: "\\cap",
    'emptyset': lambda elem: "\\emptyset",
    'rightarrow': lambda elem: "\\rightarrow",
    'leftarrow': lambda elem: "\\leftarrow",
    'leftrightarrow': lambda elem: "\\leftrightarrow",
    'Rightarrow': lambda elem: "\\Rightarrow",
    'Leftarrow': lambda elem: "\\Leftarrow",
    'Leftrightarrow': lambda elem: "\\Leftrightarrow",
    'equiv': lambda elem: "\\equiv",
    'approx': lambda elem: "\\approx",
    'sim': lambda elem: "\\sim",
    'propto': lambda elem: "\\propto",
    'langle': lambda elem: "\\langle",
    'rangle': lambda elem: "\\rangle",
    'lceil': lambda elem: "\\lceil",
    'rceil': lambda elem: "\\rceil",
    'lfloor': lambda elem: "\\lfloor",
    'rfloor': lambda elem: "\\rfloor",
    'vec': lambda elem: f"\\vec{{{get_omml_text(elem)}}}",
    'bar': lambda elem: f"\\overline{{{get_omml_text(elem)}}}",
    'hat': lambda elem: f"\\hat{{{get_omml_text(elem)}}}",
    'tilde': lambda elem: f"\\tilde{{{get_omml_text(elem)}}}",
}

def get_n():
    """Get the OMML namespace"""
    return "{http://schemas.openxmlformats.org/officeDocument/2006/math}"

n = get_n()

def get_omml_text(elem):
    """Extract text from OMML element"""
    if elem is None:
        return ""
    text = ""
    for t in elem.iter(f"{n}t"):
        if t.text:
            text += t.text
    return text

def omml_to_latex(omml_xml):
    """Convert OMML (Office Math Markup Language) to LaTeX"""
    if not omml_xml:
        return None
    
    try:
        # Parse the OMML XML
        if isinstance(omml_xml, str):
            root = ET.fromstring(omml_xml)
        else:
            root = omml_xml
        
        latex = convert_omml_element(root)
        return latex
    except Exception as e:
        return f"<!-- OMML conversion error: {e} -->"

def convert_omml_element(elem):
    """Recursively convert OMML element to LaTeX"""
    if elem is None:
        return ""
    
    # Remove namespace for easier processing
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    
    result = ""
    
    # Handle different OMML elements
    if tag == 't':  # Regular text
        return elem.text or ""
    
    elif tag == 'r':  # Run
        for child in elem:
            result += convert_omml_element(child)
        return result
    
    elif tag == 'i':  # Italic
        text = ""
        for child in elem:
            text += convert_omml_element(child)
        return f"{{{text}}}" if text else ""
    
    elif tag == 'num' or tag == 'den':  # Fraction numerator/denominator
        text = ""
        for child in elem:
            text += convert_omml_element(child)
        return text
    
    elif tag == 'f':  # Fraction
        num = elem.find(f"{n}num")
        den = elem.find(f"{n}den")
        num_text = convert_omml_element(num) if num is not None else ""
        den_text = convert_omml_element(den) if den is not None else ""
        return f"\\frac{{{num_text}}}{{{den_text}}}"
    
    elif tag == 'rad':  # Radical/Square root
        deg = elem.find(f"{n}deg")
        e = elem.find(f"{n}e")
        e_text = convert_omml_element(e) if e is not None else ""
        if deg is not None and get_omml_text(deg):
            deg_text = get_omml_text(deg)
            return f"\\sqrt[{deg_text}]{{{e_text}}}"
        return f"\\sqrt{{{e_text}}}"
    
    elif tag == 'sSub':  # Subscript
        e = elem.find(f"{n}e")
        sub = elem.find(f"{n}sub")
        e_text = convert_omml_element(e) if e is not None else ""
        sub_text = convert_omml_element(sub) if sub is not None else ""
        return f"{e_text}_{{{sub_text}}}"
    
    elif tag == 'sSup':  # Superscript
        e = elem.find(f"{n}e")
        sup = elem.find(f"{n}sup")
        e_text = convert_omml_element(e) if e is not None else ""
        sup_text = convert_omml_element(sup) if sup is not None else ""
        return f"{e_text}^{{{sup_text}}}"
    
    elif tag == 'sSubSup':  # Subscript and Superscript
        e = elem.find(f"{n}e")
        sub = elem.find(f"{n}sub")
        sup = elem.find(f"{n}sup")
        e_text = convert_omml_element(e) if e is not None else ""
        sub_text = convert_omml_element(sub) if sub is not None else ""
        sup_text = convert_omml_element(sup) if sup is not None else ""
        return f"{e_text}_{{{sub_text}}}^{{{sup_text}}}"
    
    elif tag == 'nary':  # N-ary operators (sum, integral, etc.)
        sub = elem.find(f"{n}sub")
        sup = elem.find(f"{n}sup")
        e = elem.find(f"{n}e")
        
        # Get the operator
        op_elem = elem.find(f"{n}naryPr/{n}chr")
        op = op_elem.get(f"{n}val", "") if op_elem is not None else ""
        
        sub_text = convert_omml_element(sub) if sub is not None else ""
        sup_text = convert_omml_element(sup) if sup is not None else ""
        e_text = convert_omml_element(e) if e is not None else ""
        
        op_map = {
            '\u2211': '\\sum',
            '\u220F': '\\prod',
            '\u222B': '\\int',
            '\u222C': '\\iint',
            '\u222D': '\\iiint',
            '\u222E': '\\oint',
        }
        
        latex_op = op_map.get(op, op)
        
        limits = ""
        if sub_text or sup_text:
            limits = f"_{{{sub_text}}}^{{{sup_text}}}" if sub_text and sup_text else (f"_{{{sub_text}}}" if sub_text else f"^{{{sup_text}}}")
        
        return f"{latex_op}{limits} {e_text}"
    
    elif tag == 'lim':  # Limit
        e = elem.find(f"{n}e")
        sub = elem.find(f"{n}sub")
        e_text = convert_omml_element(e) if e is not None else ""
        sub_text = convert_omml_element(sub) if sub is not None else ""
        return f"\\lim_{{{sub_text}}} {e_text}"
    
    elif tag == 'limLow':  # Limit with underscript
        e = elem.find(f"{n}e")
        sub = elem.find(f"{n}lim")
        e_text = convert_omml_element(e) if e is not None else ""
        sub_text = convert_omml_element(sub) if sub is not None else ""
        return f"{e_text}_{{{sub_text}}}"
    
    elif tag == 'limUpp':  # Limit with overscript
        e = elem.find(f"{n}e")
        sup = elem.find(f"{n}lim")
        e_text = convert_omml_element(e) if e is not None else ""
        sup_text = convert_omml_element(sup) if sup is not None else ""
        return f"{e_text}^{{{sup_text}}}"
    
    elif tag == 'acc':  # Accent
        e = elem.find(f"{n}e")
        acc_pr = elem.find(f"{n}accPr")
        chr_elem = acc_pr.find(f"{n}chr") if acc_pr is not None else None
        chr_val = chr_elem.get(f"{n}val", "^") if chr_elem is not None else "^"
        
        e_text = convert_omml_element(e) if e is not None else ""
        
        accent_map = {
            '^': '\\hat',
            '~': '\\tilde',
            '\u00AF': '\\overline',
            '\u00B7': '\\dot',
            '\u00A8': '\\ddot',
            '\u2192': '\\vec',
        }
        
        latex_accent = accent_map.get(chr_val, '\\hat')
        return f"{latex_accent}{{{e_text}}}"
    
    elif tag == 'bar':  # Bar accent
        e = elem.find(f"{n}e")
        e_text = convert_omml_element(e) if e is not None else ""
        return f"\\overline{{{e_text}}}"
    
    elif tag == 'eqArr':  # Equation array
        rows = []
        for row in elem.findall(f"{n}e"):
            row_text = convert_omml_element(row)
            rows.append(row_text)
        return "\\begin{aligned}\n" + "\\\\\n".join(rows) + "\n\\end{aligned}"
    
    elif tag == 'm':  # Matrix
        rows = []
        for mr in elem.findall(f"{n}mr"):
            cells = []
            for mc in mr.findall(f"{n}e"):
                cells.append(convert_omml_element(mc))
            rows.append(" & ".join(cells))
        return "\\begin{pmatrix}\n" + " \\\\\n".join(rows) + "\n\\end{pmatrix}"
    
    elif tag == 'e':  # Generic element
        for child in elem:
            result += convert_omml_element(child)
        return result
    
    elif tag == 'rPr':  # Run properties - skip
        return ""
    
    elif tag == 'sty':  # Style - skip
        return ""
    
    elif tag == 'phant':  # Phantom
        e = elem.find(f"{n}e")
        return convert_omml_element(e) if e is not None else ""
    
    elif tag == 'd':  # Delimiter
        e = elem.find(f"{n}e")
        d_pr = elem.find(f"{n}dPr")
        beg = d_pr.get(f"{n}beg", "") if d_pr is not None else ""
        end = d_pr.get(f"{n}end", "") if d_pr is not None else ""
        e_text = convert_omml_element(e) if e is not None else ""
        
        delim_map = {
            '(': '(',
            ')': ')',
            '[': '[',
            ']': ']',
            '{': '\\{',
            '}': '\\}',
            '|': '|',
            '\u27E8': '\\langle',
            '\u27E9': '\\rangle',
        }
        
        beg_latex = delim_map.get(beg, beg)
        end_latex = delim_map.get(end, end)
        return f"{beg_latex}{e_text}{end_latex}"
    
    elif tag == 'groupChr':  # Group character
        e = elem.find(f"{n}e")
        return convert_omml_element(e) if e is not None else ""
    
    elif tag == 'oMath':  # Math paragraph
        for child in elem:
            result += convert_omml_element(child)
        return result
    
    elif tag == 'oMathPara':  # Math paragraph wrapper
        for child in elem:
            result += convert_omml_element(child)
        return result
    
    else:
        # For unknown tags, try to extract text from children
        for child in elem:
            result += convert_omml_element(child)
        # Also get any direct text
        if elem.text:
            result = elem.text + result
        return result

def extract_math_from_paragraph(para):
    """Extract mathematical formulas from a paragraph"""
    formulas = []
    
    # Check for OMML math elements
    for math_elem in para._element.findall(f".//{n}oMath"):
        latex = convert_omml_element(math_elem)
        if latex and latex.strip():
            formulas.append({
                'type': 'inline_math',
                'latex': latex,
                'omml': ET.tostring(math_elem, encoding='unicode')
            })
    
    # Check for display math (separate math paragraphs)
    for math_para in para._element.findall(f".//{n}oMathPara"):
        for math_elem in math_para.findall(f".//{n}oMath"):
            latex = convert_omml_element(math_elem)
            if latex and latex.strip():
                formulas.append({
                    'type': 'display_math',
                    'latex': latex,
                    'omml': ET.tostring(math_elem, encoding='unicode')
                })
    
    return formulas

def extract_images_from_docx(doc, output_dir):
    """Extract all images from the DOCX file"""
    images = []
    os.makedirs(output_dir, exist_ok=True)
    
    # Get the document part
    document_part = doc.part
    
    # Track image relationships
    image_map = {}
    
    # Extract images from document relationships
    for rel_id, rel in document_part.rels.items():
        if "image" in rel.reltype:
            image_part = rel.target_part
            image_data = image_part.blob
            image_ext = os.path.splitext(rel.target_part.partname)[1].lstrip('.')
            if not image_ext:
                image_ext = 'png'
            
            # Generate filename
            image_name = f"image_{len(image_map) + 1}.{image_ext}"
            image_path = os.path.join(output_dir, image_name)
            
            # Save image
            with open(image_path, 'wb') as f:
                f.write(image_data)
            
            image_map[rel.rId] = {
                'path': image_path,
                'name': image_name,
                'content_type': image_part.content_type,
                'size': len(image_data)
            }
    
    # Now find images in the document and associate with context
    img_counter = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                for blip in blips:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id and embed_id in image_map:
                        img_info = image_map[embed_id].copy()
                        img_counter += 1
                        
                        # Try to find caption (next paragraph or same paragraph text)
                        caption = para.text.strip() if para.text.strip() else "No caption"
                        
                        # Get image dimensions if available
                        extent = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}extent')
                        width = height = None
                        if extent is not None:
                            width = int(extent.get('cx', 0)) / 9525  # Convert EMU to pixels
                            height = int(extent.get('cy', 0)) / 9525
                        
                        img_info.update({
                            'index': img_counter,
                            'context': caption[:200],  # First 200 chars as context
                            'paragraph_index': doc.paragraphs.index(para) if para in doc.paragraphs else -1,
                            'width_px': width,
                            'height_px': height,
                            'section': get_section_for_paragraph(para, doc)
                        })
                        images.append(img_info)
    
    # Also check for inline shapes (simpler image embedding)
    for para in doc.paragraphs:
        for inline_shape in para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
            imagedata = inline_shape.findall('.//{urn:schemas-microsoft-com:vml}imagedata')
            for img_data in imagedata:
                embed_id = img_data.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}relid')
                if embed_id and embed_id in image_map:
                    img_info = image_map[embed_id].copy()
                    img_counter += 1
                    
                    caption = para.text.strip() if para.text.strip() else "No caption"
                    
                    img_info.update({
                        'index': img_counter,
                        'context': caption[:200],
                        'paragraph_index': doc.paragraphs.index(para) if para in doc.paragraphs else -1,
                        'section': get_section_for_paragraph(para, doc)
                    })
                    images.append(img_info)
    
    return images

def get_section_for_paragraph(para, doc):
    """Determine which section a paragraph belongs to based on heading styles"""
    current_section = "Document Body"
    
    para_idx = doc.paragraphs.index(para) if para in doc.paragraphs else -1
    if para_idx < 0:
        return current_section
    
    # Look backwards for the last heading
    for i in range(para_idx, -1, -1):
        prev_para = doc.paragraphs[i]
        style_name = prev_para.style.name if prev_para.style else ""
        
        if style_name and ('Heading' in style_name or style_name.startswith('Heading')):
            current_section = prev_para.text.strip()
            break
    
    return current_section

def extract_document_structure(doc):
    """Extract the full document structure with paragraphs, formulas, and context"""
    structure = []
    current_section = "Document Body"
    para_index = 0
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()
        
        # Check if this is a heading
        is_heading = 'Heading' in style_name or style_name.startswith('Heading')
        if is_heading and text:
            current_section = text
            level = int(style_name.replace('Heading', '').strip()) if style_name.replace('Heading', '').strip().isdigit() else 1
        
        # Extract formulas from this paragraph
        formulas = extract_math_from_paragraph(para)
        
        # Determine paragraph type
        para_type = "normal"
        if is_heading:
            para_type = "heading"
        elif para.alignment:
            if para.alignment == 2:  # Center
                para_type = "centered"
        
        para_info = {
            'index': para_index,
            'section': current_section,
            'style': style_name,
            'type': para_type,
            'text': text,
            'formulas': formulas,
            'has_formulas': len(formulas) > 0
        }
        
        structure.append(para_info)
        para_index += 1
    
    return structure

def extract_tables(doc):
    """Extract tables from the document"""
    tables = []
    
    for tbl in doc.tables:
        table_data = {
            'rows': [],
            'num_rows': tbl.rows.__len__(),
            'num_cols': 0
        }
        
        for row in tbl.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
                if len(row_data) > table_data['num_cols']:
                    table_data['num_cols'] = len(row_data)
            table_data['rows'].append(row_data)
        
        tables.append(table_data)
    
    return tables

def main():
    print("=" * 80)
    print("DOCX Content Extractor for Typst Conversion")
    print("=" * 80)
    print(f"\nProcessing: {DOCX_PATH}\n")
    
    # Verify file exists
    if not os.path.exists(DOCX_PATH):
        print(f"ERROR: File not found: {DOCX_PATH}")
        return
    
    # Load document
    doc = Document(DOCX_PATH)
    
    # Create output directory
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # Extract document structure
    print("Extracting document structure...")
    structure = extract_document_structure(doc)
    
    # Extract images
    print("Extracting images...")
    images_dir = os.path.join(OUTPUT_DIR, "images")
    images = extract_images_from_docx(doc, images_dir)
    
    # Extract tables
    print("Extracting tables...")
    tables = extract_tables(doc)
    
    # Compile results
    all_formulas = []
    for para in structure:
        for formula in para['formulas']:
            all_formulas.append({
                'paragraph_index': para['index'],
                'section': para['section'],
                'context': para['text'][:100] if para['text'] else "",
                'type': formula['type'],
                'latex': formula['latex']
            })
    
    # Generate report
    report = []
    report.append("=" * 80)
    report.append("EXTRACTION REPORT")
    report.append("=" * 80)
    
    report.append(f"\nDocument: {os.path.basename(DOCX_PATH)}")
    report.append(f"Total paragraphs: {len(structure)}")
    report.append(f"Total formulas found: {len(all_formulas)}")
    report.append(f"Total images found: {len(images)}")
    report.append(f"Total tables found: {len(tables)}")
    
    # Formulas section
    report.append("\n" + "=" * 80)
    report.append("MATHEMATICAL FORMULAS")
    report.append("=" * 80)
    
    if all_formulas:
        for i, formula in enumerate(all_formulas, 1):
            report.append(f"\n--- Formula #{i} ---")
            report.append(f"Section: {formula['section']}")
            report.append(f"Paragraph Index: {formula['paragraph_index']}")
            report.append(f"Type: {formula['type']}")
            report.append(f"Context: {formula['context']}")
            report.append(f"LaTeX: {formula['latex']}")
    else:
        report.append("\nNo formulas found in the document.")
    
    # Images section
    report.append("\n" + "=" * 80)
    report.append("IMAGES")
    report.append("=" * 80)
    
    if images:
        for img in images:
            report.append(f"\n--- Image #{img['index']} ---")
            report.append(f"File: {img['name']}")
            report.append(f"Path: {img['path']}")
            report.append(f"Section: {img['section']}")
            report.append(f"Context/Caption: {img['context']}")
            report.append(f"Size: {img.get('width_px', 'N/A')} x {img.get('height_px', 'N/A')} px")
            report.append(f"Content Type: {img['content_type']}")
    else:
        report.append("\nNo images found in the document.")
    
    # Tables section
    report.append("\n" + "=" * 80)
    report.append("TABLES")
    report.append("=" * 80)
    
    if tables:
        for i, table in enumerate(tables, 1):
            report.append(f"\n--- Table #{i} ---")
            report.append(f"Dimensions: {table['num_rows']} rows x {table['num_cols']} columns")
            for row_idx, row in enumerate(table['rows']):
                report.append(f"  Row {row_idx + 1}: {row}")
    else:
        report.append("\nNo tables found in the document.")
    
    # Document structure summary
    report.append("\n" + "=" * 80)
    report.append("DOCUMENT STRUCTURE SUMMARY")
    report.append("=" * 80)
    
    sections = {}
    for para in structure:
        section = para['section']
        if section not in sections:
            sections[section] = {'paragraphs': 0, 'formulas': 0, 'has_images': False}
        sections[section]['paragraphs'] += 1
        sections[section]['formulas'] += len(para['formulas'])
    
    for section, stats in sections.items():
        report.append(f"\nSection: {section}")
        report.append(f"   Paragraphs: {stats['paragraphs']}")
        report.append(f"   Formulas: {stats['formulas']}")
    
    # Print report
    report_text = "\n".join(report)
    print("\n" + report_text)
    
    # Save report to file
    report_path = os.path.join(OUTPUT_DIR, "extraction_report.txt")
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(report_text)
    print(f"\n\nReport saved to: {report_path}")
    
    # Save formulas to separate file for easy import
    formulas_path = os.path.join(OUTPUT_DIR, "formulas.txt")
    with open(formulas_path, 'w', encoding='utf-8') as f:
        f.write("# Extracted Formulas (LaTeX format)\n")
        f.write("# Format: [Section] Paragraph #index: LaTeX\n\n")
        for i, formula in enumerate(all_formulas, 1):
            f.write(f"# Formula {i}\n")
            f.write(f"[{formula['section']}]\n")
            f.write(f"Paragraph #{formula['paragraph_index']}\n")
            f.write(f"Context: {formula['context']}\n")
            f.write(f"{formula['latex']}\n\n")
    print(f"Formulas saved to: {formulas_path}")
    
    # Save JSON data for programmatic access
    import json
    json_path = os.path.join(OUTPUT_DIR, "extraction_data.json")
    json_data = {
        'document': os.path.basename(DOCX_PATH),
        'total_paragraphs': len(structure),
        'total_formulas': len(all_formulas),
        'total_images': len(images),
        'total_tables': len(tables),
        'formulas': all_formulas,
        'images': images,
        'tables': tables,
        'structure': structure
    }
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, indent=2, ensure_ascii=False)
    print(f"JSON data saved to: {json_path}")
    
    print("\n" + "=" * 80)
    print("Extraction complete!")
    print(f"Output directory: {OUTPUT_DIR}")
    print("=" * 80)

if __name__ == "__main__":
    main()
